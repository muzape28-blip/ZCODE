# python-docx — bikin file Word beneran dari HP (TESTED di ZCODE)
# Butuh: install python-docx di INSTALL MODULES
# Hasil: laporan.docx di workspace — buka via file manager / kirim WA

from docx import Document

doc = Document()
doc.add_heading("Laporan Praktikum", 0)
doc.add_paragraph("Dibuat langsung dari HP pakai ZCODE.")

t = doc.add_table(rows=1, cols=2)
t.style = "Light Grid Accent 1"
t.rows[0].cells[0].text = "Nama"
t.rows[0].cells[1].text = "Nilai"
for nama, nilai in [("Ani", 85), ("Budi", 92)]:
    row = t.add_row()
    row.cells[0].text = nama
    row.cells[1].text = str(nilai)

doc.save("laporan.docx")
print("Tersimpan: laporan.docx — cek file manager!")
